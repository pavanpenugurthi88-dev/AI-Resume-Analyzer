"""
File Extractor Service
Handles PDF and DOCX text extraction
"""

import fitz  # PyMuPDF
import docx
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text("text")
        doc.close()
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise ValueError(f"Could not extract text from PDF: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        doc = docx.Document(file_path)
        text_parts = []

        # Paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())

        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise ValueError(f"Could not extract text from DOCX: {e}")


def extract_text(file_path: str) -> str:
    """Extract text from file based on extension."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")


def validate_file(file_path: str, max_size_mb: int = 10) -> bool:
    """Validate file size and type."""
    file_size = os.path.getsize(file_path) / (1024 * 1024)  # Convert to MB
    if file_size > max_size_mb:
        raise ValueError(f"File size ({file_size:.1f}MB) exceeds maximum allowed ({max_size_mb}MB)")

    ext = Path(file_path).suffix.lower()
    if ext not in [".pdf", ".docx", ".doc"]:
        raise ValueError(f"Unsupported file type: {ext}")

    return True
